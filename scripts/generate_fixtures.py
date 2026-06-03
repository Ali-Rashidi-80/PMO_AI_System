#!/usr/bin/env python3
"""Generate test fixture documents."""

from __future__ import annotations

from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
FIX = ROOT / "tests" / "fixtures" / "documents"
PMO = ROOT / "pmo_docs"


def write_text_fixtures():
    FIX.mkdir(parents=True, exist_ok=True)
    (FIX / "weekly_reports").mkdir(exist_ok=True)

    texts = {
        "contract_full_fa.txt": (
            "قرارداد پیمانکاری — بند ۱۲: تأخیر تحویل بیش از ۳۰ روز مشمول جریمه ۰.۱٪ روزانه.\n"
            "بند ۱۵: Force Majeure فقط با تأیید کارفرما.\n" * 2
        ),
        "contract_mixed_fa_en.txt": "Contract clause 12 — تأخیر فاز ۳ Project Alpha delay penalty.\n" * 3,
        "weekly_reports/weekly_report_delay.txt": (
            "گزارش هفتگی: تأخیر فاز ۳ به دلیل Claim پیمانکار.\n" * 4
        ),
        "weekly_reports/weekly_report_material.txt": (
            "گزارش: کمبود مصالح فولادی در سایت.\n" * 4
        ),
        "weekly_reports/weekly_report_stakeholder.txt": (
            "گزارش: نارضایتی ذینفع محلی از سر و صدا.\n" * 4
        ),
        "empty.txt": "",
        "corrupt_binary.txt": "ok prefix \x00\x01\xff " + "تست " * 3,
        "large_50kb.txt": "بند قرارداد تأخیر. " * 2500,
        "sample.md": "# گزارش\n\n## تأخیر\n\nفاز ۳ با ۱۵ روز تأخیر.\n" * 3,
    }
    for name, content in texts.items():
        p = FIX / name
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content, encoding="utf-8")

    try:
        from docx import Document

        doc = Document()
        doc.add_paragraph("بند قرارداد: تأخیر تحویل مشمول جریمه است.")
        doc.save(FIX / "contract_clause.docx")
        doc2 = Document()
        doc2.add_paragraph("گزارش هفتگی: Claim فاز ۳ ثبت شد.")
        doc2.save(FIX / "weekly_reports" / "weekly_report.docx")
        bad = FIX / "invalid.docx"
        bad.write_bytes(b"not a zip")
    except ImportError:
        pass

    try:
        from pypdf import PdfWriter

        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        with open(FIX / "contract_summary.pdf", "wb") as f:
            writer.write(f)
        (FIX / "scan_only.pdf").write_bytes(
            b"%PDF-1.4 minimal\n"  # invalid/minimal — extract may fail
        )
    except ImportError:
        pass

    oracle = ROOT / "tests" / "fixtures" / "expected"
    oracle.mkdir(parents=True, exist_ok=True)
    (oracle / "annotated_risks.json").write_text(
        """{
  "project_risks": [
    {"risk_title": "تأخیر فاز ۳", "severity": "High", "evidence": "Claim", "recommended_action": "اخطار"},
    {"risk_title": "کمبود مصالح", "severity": "Medium", "evidence": "گزارش", "recommended_action": "تأمین"},
    {"risk_title": "ذینفع", "severity": "Low", "evidence": "نارضایتی", "recommended_action": "جلسه"}
  ]
}""",
        encoding="utf-8",
    )

    # Copy subset to pmo_docs for demo
    PMO.mkdir(exist_ok=True)
    (PMO / "weekly_reports").mkdir(exist_ok=True)
    for src in [
        "contract_full_fa.txt",
        "weekly_reports/weekly_report_delay.txt",
        "weekly_reports/weekly_report_material.txt",
    ]:
        sp = FIX / src
        if sp.is_file():
            dp = PMO / src
            dp.parent.mkdir(parents=True, exist_ok=True)
            dp.write_text(sp.read_text(encoding="utf-8"), encoding="utf-8")

    print(f"Fixtures written to {FIX}")


if __name__ == "__main__":
    write_text_fixtures()
